from docx import Document
import os
# 将一个文件夹的word文件转换称txt文件,01/23 修改代码，适配win
path = r'C:\roczhang\WHPU\zen\政策文本'
path_txt = r'C:\roczhang\WHPU\zen\data'  # 这个文件夹用来装txt文件
file_list = os.listdir(path)  # 读取出docx文件夹所有的文件名字

for file in file_list:
    file_path = path + '/' + file
    # print(file_path)
    doc = Document(file_path)  # 读取docx文件
    f = open(path_txt + '/' + file.split('.')[0] + '.txt', 'a')  # 把.docx的后缀改成txt，并创建txt文件。
    print(file_path + '\n')
    for paragraph in doc.paragraphs:
        f.write(paragraph.text)  # 将docx段落写入txt文件
        print(paragraph.text + '\n')
    f.close()  # txt文件使用完成后关闭

# 更好的写法
for file in file_list:
    file_path = path + '/' + file
    # print(file_path)
    doc = Document(file_path)  # 读取docx文件
    with open(path_txt + '/' + file.split('.')[0] + '.txt', 'a', encoding='utf-8') as f:
        for paragraph in doc.paragraphs:
            f.write(paragraph.text)  # 将docx段落写入txt文件
            print(paragraph.text + '\n')

